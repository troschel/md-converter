#!/usr/bin/env python3
"""
Converts .mdx and .json files to .docx and back

Usage:
    # Source  -> .docx
    python convert_for_translation.py to-docx   ./source   ./for_translation

    # .docx  -> .mdx/.json
    python convert_for_translation.py from-docx ./translated ./ergebnis

    change ./source or ./translated with the current directory where your files are

"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

SOURCE_EXTENSIONS = {".mdx", ".json"}
MONO_FONT = "Consolas"


def set_monospace(document: Document) -> None:
    """Set standard font of document to Monospace"""
    style = document.styles["Normal"]
    style.font.name = MONO_FONT
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), MONO_FONT)


def file_to_docx(src: Path, dst: Path) -> None:
    text = src.read_text(encoding="utf-8")
    document = Document()
    set_monospace(document)
    # removes linespace, every line becomes paragraph
    lines = text.split("\n")
    for line in lines:
        document.add_paragraph(line)
    dst.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(dst))


def docx_to_file(src: Path, dst: Path) -> None:
    document = Document(str(src))
    lines = [p.text for p in document.paragraphs]
    text = "\n".join(lines)
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text(text, encoding="utf-8")


def run_to_docx(in_dir: Path, out_dir: Path) -> int:
    count = 0
    for src in sorted(in_dir.rglob("*")):
        if src.is_file() and src.suffix.lower() in SOURCE_EXTENSIONS:
            rel = src.relative_to(in_dir)
            dst = out_dir / rel.with_name(rel.name + ".docx")
            file_to_docx(src, dst)
            print(f"  {rel}  ->  {dst.relative_to(out_dir)}")
            count += 1
    return count


def run_from_docx(in_dir: Path, out_dir: Path) -> int:
    count = 0
    for src in sorted(in_dir.rglob("*.docx")):
        if src.is_file():
            rel = src.relative_to(in_dir)
            # "seite.mdx.docx" -> "seite.mdx"
            original_name = rel.name[: -len(".docx")]
            dst = out_dir / rel.with_name(original_name)
            if dst.suffix.lower() not in SOURCE_EXTENSIONS:
                print(f"  uebersprungen (kein .mdx/.json): {rel}")
                continue
            docx_to_file(src, dst)
            print(f"  {rel}  ->  {dst.relative_to(out_dir)}")
            count += 1
    return count


def main() -> None:
    if len(sys.argv) != 4 or sys.argv[1] not in {"to-docx", "from-docx"}:
        print(__doc__)
        sys.exit(1)

    mode, in_dir, out_dir = sys.argv[1], Path(sys.argv[2]), Path(sys.argv[3])
    if not in_dir.is_dir():
        print(f"Fehler: '{in_dir}' ist kein Ordner.")
        sys.exit(1)

    if mode == "to-docx":
        print(f"Konvertiere .mdx/.json aus '{in_dir}' nach .docx in '{out_dir}':")
        n = run_to_docx(in_dir, out_dir)
    else:
        print(f"Konvertiere .docx aus '{in_dir}' zurueck nach .mdx/.json in '{out_dir}':")
        n = run_from_docx(in_dir, out_dir)

    print(f"\nFertig. {n} Datei(en) konvertiert.")


if __name__ == "__main__":
    main()